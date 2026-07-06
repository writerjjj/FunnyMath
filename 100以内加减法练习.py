"""生成 100 以内加减法 Word 练习题（单页写满）。

规则:
    - 数值范围 [1, 100]，参数与答案均不为 0
    - 加法 a + b <= 100（a, b >= 1）
    - 减法 a - b >= 1（a > b >= 1）
    - 一次仅生成一页，页内题目互不重复
    - 题号用圆圈符号（①②③...）
    - 题目区使用无边框表格排版，按页面可用高度均分行高，铺满整页
    - 答案作为正文段落写在题目下方
    - 页眉与页脚显示同一份生成时间戳（精确到秒）
"""

from __future__ import annotations

import argparse
import random
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


@dataclass(frozen=True)
class Problem:
    """一道题目：a op b。"""

    a: int
    op: str
    b: int

    @property
    def answer(self) -> int:
        return self.a + self.b if self.op == "+" else self.a - self.b

    def question(self) -> str:
        return f"{self.a} {self.op} {self.b} ="


def make_problem(max_value: int, rng: random.Random) -> Problem:
    """随机生成一个合法的加减法题目：参数与答案均 >= 1。"""
    op = rng.choice(["+", "-"])
    if op == "+":
        # a>=1, b>=1, a+b<=max_value -> a 至多 max_value-1
        a = rng.randint(1, max_value - 1)
        b = rng.randint(1, max_value - a)
    else:
        # a>=2, b>=1, a-b>=1 -> b<=a-1
        a = rng.randint(2, max_value)
        b = rng.randint(1, a - 1)
    return Problem(a, op, b)


def build_page_problems(
    count: int, max_value: int, rng: random.Random
) -> list[Problem]:
    """构造单页题目，题面（含运算方向）互不重复。"""
    seen: set[tuple[int, str, int]] = set()
    problems: list[Problem] = []
    # 100 以内加减法组合数远大于常见页题量，无死循环风险
    while len(problems) < count:
        p = make_problem(max_value, rng)
        key = (p.a, p.op, p.b)
        if key in seen:
            continue
        seen.add(key)
        problems.append(p)
    return problems


# 圆圈数字通过 Unicode 代码点构造，避免在源码中直接嵌入生僻字符：
#   1-20   -> U+2460..U+2473  (① .. ⑳)
#   21-35  -> U+3251..U+325F  (㉑ .. ㉟)
#   36-50  -> U+32B1..U+32BF  (㊱ .. ㊿)
# 超过 50 的题号无对应 Unicode 圆圈数字，会在 circled() 中回退成 (n)。
_CIRCLED_RANGES: list[tuple[int, int]] = [
    (0x2460, 20),  # ①-⑳
    (0x3251, 15),  # ㉑-㉟
    (0x32B1, 15),  # ㊱-㊿
]
_CIRCLED_NUMBERS: list[str] = [
    chr(start + offset)
    for start, count in _CIRCLED_RANGES
    for offset in range(count)
]


def circled(n: int) -> str:
    """返回 n 对应的圆圈数字；超出 Unicode 覆盖范围（>50）回退到 (n)。"""
    if 1 <= n <= len(_CIRCLED_NUMBERS):
        return _CIRCLED_NUMBERS[n - 1]
    return f"({n})"


def _hide_table_borders(table) -> None:
    """把整张表格的边框设为 none，让表格在视觉上不显示。"""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = borders.makeelement(qn(f"w:{edge}"), {})
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _write_timestamp_to_header_and_footer(doc: Document, timestamp_text: str) -> None:
    """把同一份时间戳写到页眉和页脚（居中，小字）。"""
    section = doc.sections[0]
    for target in (section.header, section.footer):
        para = target.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 清空既有 run，避免 python-docx 默认的空 run 干扰
        for run in list(para.runs):
            run.text = ""
        run = para.add_run(timestamp_text)
        run.font.size = Pt(9)


def _write_answers_to_body(
    doc: Document, problems: list[Problem], answers_per_row: int = 10
) -> None:
    """把答案写在题目下方：每行 answers_per_row 个，使用无边框表格保证列对齐。"""
    section = doc.sections[0]
    usable_w_cm = (
        section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    )
    col_w_cm = usable_w_cm / answers_per_row

    # 在表格前留一小段间距，代替段落 space_before（表格自己没有 space_before）
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(0)
    gap.paragraph_format.line_spacing = 1.0
    gap_run = gap.add_run("")
    gap_run.font.size = Pt(6)

    rows = (len(problems) + answers_per_row - 1) // answers_per_row
    table = doc.add_table(rows=rows, cols=answers_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _hide_table_borders(table)

    for r in range(rows):
        for c in range(answers_per_row):
            idx = r * answers_per_row + c
            cell = table.rows[r].cells[c]
            cell.width = Cm(col_w_cm)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            if idx < len(problems):
                run = para.add_run(f"{circled(idx + 1)}{problems[idx].answer}")
                run.font.size = Pt(13)
            else:
                para.add_run(" ")


def render_page(
    doc: Document,
    problems: list[Problem],
    columns: int,
    answer_reserve_cm: float,
) -> None:
    """在文档中写入一页题目：题目均分行高铺满正文，答案作为正文段落写在其后。"""
    section = doc.sections[0]
    usable_h_cm = (
        section.page_height.cm
        - section.top_margin.cm
        - section.bottom_margin.cm
        - answer_reserve_cm
    )
    rows = (len(problems) + columns - 1) // columns
    row_h_cm = max(0.9, usable_h_cm / rows)
    col_w_cm = (
        section.page_width.cm
        - section.left_margin.cm
        - section.right_margin.cm
    ) / columns

    table = doc.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _hide_table_borders(table)

    for r in range(rows):
        row = table.rows[r]
        row.height = Cm(row_h_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for c in range(columns):
            idx = r * columns + c
            cell = row.cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(col_w_cm)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if idx < len(problems):
                run = paragraph.add_run(
                    f"{circled(idx + 1)} {problems[idx].question()}"
                )
                run.font.size = Pt(13)
            else:
                paragraph.add_run(" ")

    _write_answers_to_body(doc, problems)


def build_document(
    per_page: int,
    columns: int,
    max_value: int,
    seed: int | None,
    output: Path,
) -> Path:
    rng = random.Random(seed)
    doc = Document()

    for section in doc.sections:
        # 显式使用 A4 纸张（21 × 29.7 cm）
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)

    # 生成时间只取一次，保证页眉与页脚上的时间完全一致
    timestamp_text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    _write_timestamp_to_header_and_footer(doc, timestamp_text)

    problems = build_page_problems(per_page, max_value, rng)
    # 答案改为 10 列 × 5 行的对齐表格，预留 ~4.5cm 保证不溢出到第二页
    render_page(doc, problems=problems, columns=columns, answer_reserve_cm=4.5)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="生成 100 以内加减法练习题 Word 文档（单页写满）"
    )
    parser.add_argument("--per-page", type=int, default=50, help="每页题量（默认 50，编号 ①–㊿ 完整覆盖）")
    parser.add_argument("--columns", type=int, default=4, help="每页列数（默认 4）")
    parser.add_argument(
        "--max", dest="max_value", type=int, default=100, help="数值上限（默认 100）"
    )
    parser.add_argument("--seed", type=int, default=None, help="随机种子（可选，便于复现）")
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("100以内加减法练习.docx"),
        help="输出文件路径",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.per_page > len(_CIRCLED_NUMBERS):
        print(
            f"提示：每页题量 {args.per_page} 超过圆圈编号支持范围"
            f"（{len(_CIRCLED_NUMBERS)}），多出部分将回退为 (n)。"
        )
    if args.per_page % args.columns != 0:
        print(
            f"提示：每页 {args.per_page} 题不能被 {args.columns} 列整除，最后一行会有空位。"
        )
    output = build_document(
        per_page=args.per_page,
        columns=args.columns,
        max_value=args.max_value,
        seed=args.seed,
        output=args.output,
    )
    print(f"已生成：{output.resolve()}")


if __name__ == "__main__":
    main()
